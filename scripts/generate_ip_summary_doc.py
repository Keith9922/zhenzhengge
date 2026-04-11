from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(31, 78, 121)
TEXT = RGBColor(44, 62, 80)
MUTED = RGBColor(107, 114, 128)
QUOTE_BG = "F3F6FA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_border(paragraph, border_color: str = "D9E3F0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), border_color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def set_east_asia_font(style_or_run, font_name: str) -> None:
    r_pr = style_or_run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def clean_md_text(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("\\", "")
    return text.strip()


def add_inline_runs(paragraph, text: str) -> None:
    text = clean_md_text(text)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10.5)
            run.font.color.rgb = ACCENT
            set_east_asia_font(run, "PingFang SC")
        else:
            tokens = re.split(r"(\*\*.*?\*\*)", part)
            for token in tokens:
                if not token:
                    continue
                if token.startswith("**") and token.endswith("**"):
                    run = paragraph.add_run(token[2:-2])
                    run.bold = True
                else:
                    run = paragraph.add_run(token)
                run.font.name = "Aptos"
                run.font.size = Pt(11)
                run.font.color.rgb = TEXT
                set_east_asia_font(run, "PingFang SC")


def ensure_style(document: Document, name: str, base: str | None = None):
    if name in document.styles:
        return document.styles[name]
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        style.base_style = document.styles[base]
    return style


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    set_east_asia_font(normal, "PingFang SC")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for level, size in [(1, 18), (2, 15), (3, 12.5), (4, 11.5)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos"
        style.font.color.rgb = ACCENT
        style.font.bold = True
        style.font.size = Pt(size)
        set_east_asia_font(style, "PingFang SC")
        style.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    list_bullet = document.styles["List Bullet"]
    list_bullet.font.name = "Aptos"
    list_bullet.font.size = Pt(11)
    set_east_asia_font(list_bullet, "PingFang SC")

    list_number = document.styles["List Number"]
    list_number.font.name = "Aptos"
    list_number.font.size = Pt(11)
    set_east_asia_font(list_number, "PingFang SC")

    quote_style = ensure_style(document, "Custom Quote", "Normal")
    quote_style.font.italic = True
    quote_style.font.size = Pt(10.5)
    quote_style.font.color.rgb = MUTED
    set_east_asia_font(quote_style, "PingFang SC")
    quote_style.paragraph_format.left_indent = Cm(0.35)
    quote_style.paragraph_format.space_after = Pt(3)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("魔法黑客松江苏知识产权专场纪要")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_east_asia_font(run, "PingFang SC")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = fp.add_run("第 ")
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = MUTED
    set_east_asia_font(prefix, "PingFang SC")
    add_page_number(fp)
    suffix = fp.add_run(" 页")
    suffix.font.size = Pt(9)
    suffix.font.color.rgb = MUTED
    set_east_asia_font(suffix, "PingFang SC")


def add_cover(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(21)
    run.font.color.rgb = ACCENT
    set_east_asia_font(run, "PingFang SC")

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    run = sub.add_run("根据分享会逐字稿整理，便于选题、方案设计与后续查阅")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    set_east_asia_font(run, "PingFang SC")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    run = meta.add_run(f"整理时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    set_east_asia_font(run, "PingFang SC")

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, QUOTE_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(
        p,
        "本文件按 `赛道方向`、`专业边界`、`问答摘要`、`比赛信息` 四部分整理，适合后续直接用作选题讨论材料、需求梳理底稿和答辩准备参考。",
    )

    document.add_paragraph()


def add_quote_block(document: Document, quote_lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, QUOTE_BG)
    p = cell.paragraphs[0]
    p.style = document.styles["Custom Quote"]
    set_paragraph_border(p)
    add_inline_runs(p, " ".join(line.lstrip("> ").strip() for line in quote_lines))
    document.add_paragraph()


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    in_code_block = False
    paragraph_buffer: list[str] = []
    quote_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        paragraph_buffer = []
        if not text:
            return
        p = document.add_paragraph()
        add_inline_runs(p, text)

    def flush_quote() -> None:
        nonlocal quote_buffer
        if not quote_buffer:
            return
        add_quote_block(document, quote_buffer)
        quote_buffer = []

    for idx, line in enumerate(lines):
        stripped = line.strip()

        if idx == 0 and stripped.startswith("# "):
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            flush_quote()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = document.add_paragraph(style="No Spacing")
            run = p.add_run(line)
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)
            set_east_asia_font(run, "PingFang SC")
            continue

        if not stripped:
            flush_paragraph()
            flush_quote()
            continue

        if stripped == "---":
            flush_paragraph()
            flush_quote()
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_buffer.append(stripped)
            continue
        else:
            flush_quote()

        heading_match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            hashes, title = heading_match.groups()
            level = len(hashes) - 1
            if level == 1:
                level = 1
            elif level > 3:
                level = 4
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, title)
            if title.startswith("7.") or title.startswith("8.") or title.startswith("9."):
                p.paragraph_format.page_break_before = True
            continue

        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered_match:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            add_inline_runs(p, ordered_match.group(2))
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet_match.group(1))
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()
    flush_quote()


def build_docx(source_md: Path, output_docx: Path) -> None:
    markdown = source_md.read_text(encoding="utf-8")
    title = markdown.splitlines()[0].lstrip("# ").strip()

    document = Document()
    configure_document(document)
    add_header_footer(document)
    add_cover(document, title)
    render_markdown(document, markdown)

    document.core_properties.title = title
    document.core_properties.subject = "知识产权黑客松分享纪要"
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.comments = "由逐字稿整理生成"

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        default="docs/ip-hackathon-expert-summary.md",
        help="Source markdown file",
    )
    parser.add_argument(
        "--output",
        default="output/doc/ip-hackathon-expert-summary.docx",
        help="Output docx file",
    )
    args = parser.parse_args()

    source_md = Path(args.source)
    output_docx = Path(args.output)
    build_docx(source_md, output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
