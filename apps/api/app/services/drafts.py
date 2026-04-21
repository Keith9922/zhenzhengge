from pathlib import Path
import logging
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import Settings, settings as global_settings
from app.core.storage import SQLiteStorage
from app.schemas.drafts import DocumentDraftCreateRequest, DocumentDraftRecord, DraftStatus
from app.services.cases import CaseService
from app.services.evidence import EvidenceService
from app.services.hermes import HermesOrchestrator
from app.services.templates import DocumentTemplateService


class DocumentDraftService:
    def __init__(
        self,
        storage: SQLiteStorage,
        *,
        case_service: CaseService,
        template_service: DocumentTemplateService,
        evidence_service: EvidenceService,
        hermes: HermesOrchestrator,
        settings: Settings | None = None,
    ) -> None:
        self.storage = storage
        self.case_service = case_service
        self.template_service = template_service
        self.evidence_service = evidence_service
        self.hermes = hermes
        self.settings = settings or global_settings

    def list_drafts(
        self,
        case_id: str | None = None,
        *,
        organization_id: str | None = None,
    ) -> list[DocumentDraftRecord]:
        return self.storage.list_document_drafts(case_id=case_id, organization_id=organization_id)

    def get_draft(self, draft_id: str, *, organization_id: str | None = None) -> DocumentDraftRecord | None:
        return self.storage.get_document_draft(draft_id, organization_id=organization_id)

    def generate_draft(
        self,
        payload: DocumentDraftCreateRequest,
        *,
        organization_id: str,
        owner_user_id: str,
    ) -> DocumentDraftRecord:
        case = self.case_service.get_case(payload.case_id, organization_id=organization_id)
        if case is None:
            raise ValueError("案件不存在")

        template = self.template_service.get_template(payload.template_key)
        if template is None:
            raise ValueError("模板不存在")

        title = f"{case.title} - {template.name}"
        evidence_items = self.evidence_service.list_packs(case.case_id, organization_id=organization_id)
        evidence_context = [self._build_evidence_context(item) for item in evidence_items]
        case_context = {
            "case_id": case.case_id,
            "title": case.title,
            "brand_name": case.brand_name,
            "suspect_name": case.suspect_name,
            "platform": case.platform,
            "risk_level": case.risk_level,
            "description": case.description,
        }
        precheck = self.hermes.submit_document_workflow(
            template.template_key, case.case_id,
            case_context=case_context,
            evidence_context=evidence_context,
        )
        if precheck.status == "skipped":
            logging.getLogger(__name__).warning("文书预检跳过：%s", precheck.detail)
        llm_result = self.hermes.generate_document_draft(
            template_name=template.name,
            case_context=case_context,
            evidence_context=evidence_context,
            variables_override=payload.variables_override,
        )
        if self.settings.draft_generation_strict and llm_result.status != "ok":
            raise RuntimeError(
                "文书生成失败：严格模式已启用，未返回可信模型结果。"
                f" detail={llm_result.detail}"
            )
        evidence_lines = self._format_evidence_reference_lines(evidence_context)
        content = llm_result.content or self._render_content(
            case_title=case.title,
            brand_name=case.brand_name,
            suspect_name=case.suspect_name,
            platform=case.platform,
            risk_level=case.risk_level,
            description=case.description,
            template_name=template.name,
            extra_variables=payload.variables_override,
            evidence_lines=evidence_lines,
        )
        return self.storage.create_document_draft(
            case_id=payload.case_id,
            organization_id=organization_id,
            owner_user_id=owner_user_id,
            template_key=payload.template_key,
            title=title,
            content=content,
        )

    def submit_review(
        self, draft_id: str, comment: str, *, organization_id: str | None = None
    ) -> DocumentDraftRecord | None:
        return self.storage.update_document_draft_review(
            draft_id=draft_id,
            organization_id=organization_id,
            status=DraftStatus.submitted,
            review_comment=comment,
        )

    def update_draft(
        self,
        *,
        draft_id: str,
        organization_id: str | None = None,
        content: str,
        title: str | None = None,
    ) -> DocumentDraftRecord | None:
        cleaned = content.strip()
        if not cleaned:
            raise ValueError("草稿内容不能为空")
        return self.storage.update_document_draft_content(
            draft_id=draft_id,
            organization_id=organization_id,
            content=cleaned,
            title=title.strip() if title else None,
        )

    def approve(self, draft_id: str, comment: str, *, organization_id: str | None = None) -> DocumentDraftRecord | None:
        return self.storage.update_document_draft_review(
            draft_id=draft_id,
            organization_id=organization_id,
            status=DraftStatus.approved,
            review_comment=comment,
        )

    def reject(self, draft_id: str, comment: str, *, organization_id: str | None = None) -> DocumentDraftRecord | None:
        return self.storage.update_document_draft_review(
            draft_id=draft_id,
            organization_id=organization_id,
            status=DraftStatus.rejected,
            review_comment=comment,
        )

    def export_draft(self, draft_id: str, *, organization_id: str | None = None) -> DocumentDraftRecord | None:
        draft = self.storage.get_document_draft(draft_id, organization_id=organization_id)
        if draft is None:
            return None

        base_dir = Path(self.storage.db_path).resolve().parent if self.storage.db_path != ":memory:" else Path.cwd()
        export_dir = base_dir / "exports" / "drafts"
        export_dir.mkdir(parents=True, exist_ok=True)
        export_path = export_dir / f"{draft_id}.docx"

        doc = Document()
        self._configure_document(doc)

        title = doc.add_heading(draft.title or "文书草稿", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"案件编号：{draft.case_id}  ").bold = True
        meta.add_run(f"模板：{draft.template_key}  ")
        meta.add_run(f"导出时间：{draft.updated_at.strftime('%Y-%m-%d %H:%M:%S')}")

        lines = (draft.content or "").strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line == "---":
                doc.add_paragraph()
                continue
            # Handle headings
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif re.match(r"^\d+\.\s+", line):
                p = doc.add_paragraph(style="List Number")
                self._append_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            # Handle list items
            elif line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                self._append_inline_runs(p, line[2:])
            elif line.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                run = p.add_run(line[2:])
                run.italic = True
            else:
                p = doc.add_paragraph()
                self._append_inline_runs(p, line)

        doc.save(str(export_path))

        relative_path = export_path.relative_to(base_dir).as_posix()
        return self.storage.set_document_draft_export_path(
            draft_id,
            relative_path,
            organization_id=organization_id,
        )

    def _render_content(
        self,
        *,
        case_title: str,
        brand_name: str,
        suspect_name: str,
        platform: str,
        risk_level: str,
        description: str,
        template_name: str,
        extra_variables: dict[str, str],
        evidence_lines: str,
    ) -> str:
        extra_text = "\n".join(f"- {key}: {value}" for key, value in extra_variables.items()) or "- 无"
        return (
            f"# {template_name}\n\n"
            f"## 案件标题\n{case_title}\n\n"
            f"## 品牌对象\n{brand_name}\n\n"
            f"## 疑似主体\n{suspect_name}\n\n"
            f"## 来源平台\n{platform}\n\n"
            f"## 风险等级\n{risk_level}\n\n"
            f"## 案件说明\n{description}\n\n"
            f"## 证据依据\n{evidence_lines}\n\n"
            f"## 补充变量\n{extra_text}\n\n"
            "## 使用提示\n"
            "本草稿用于内部整理和法务审核前预览，正式对外动作仍需人工确认。\n"
        )

    def _build_evidence_context(self, item) -> dict[str, str]:
        return {
            "evidence_pack_id": item.evidence_pack_id,
            "source_title": item.source_title,
            "source_url": item.source_url,
            "note": item.note or "",
            "capture_channel": item.capture_channel,
            "captured_at": item.created_at.isoformat(),
            "hash_sha256": item.hash_sha256,
            "html_excerpt": self._extract_html_excerpt(self.evidence_service.read_html(item)),
        }

    @staticmethod
    def _format_evidence_reference_lines(items: list[dict[str, str]]) -> str:
        if not items:
            return "- 无"
        lines: list[str] = []
        for index, item in enumerate(items, start=1):
            lines.append(
                (
                    f"- {index}. EvidenceID={item.get('evidence_pack_id', '')} | "
                    f"URL={item.get('source_url', '')} | HASH={item.get('hash_sha256', '')} | "
                    f"取证时间={item.get('captured_at', '')}"
                )
            )
        return "\n".join(lines)

    @staticmethod
    def _extract_html_excerpt(raw_html: str, limit: int = 160) -> str:
        text = re.sub(r"<script[\s\S]*?</script>", " ", raw_html, flags=re.IGNORECASE)
        text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        if len(text) <= limit:
            return text
        return f"{text[:limit].rstrip()}..."

    @staticmethod
    def _append_inline_runs(paragraph, text: str) -> None:
        parts = re.split(r"\*\*(.+?)\*\*", text)
        for index, part in enumerate(parts):
            run = paragraph.add_run(part)
            if index % 2 == 1:
                run.bold = True

    @staticmethod
    def _configure_document(doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal_style = doc.styles["Normal"]
        DocumentDraftService._set_style_font(normal_style, "宋体", 11)

        heading1 = doc.styles["Heading 1"]
        DocumentDraftService._set_style_font(heading1, "黑体", 14)

        heading2 = doc.styles["Heading 2"]
        DocumentDraftService._set_style_font(heading2, "黑体", 12)

    @staticmethod
    def _set_style_font(style, font_name: str, size: int) -> None:
        style.font.name = font_name
        style.font.size = Pt(size)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
